# app.py

import os
import re
import io
import zipfile
import tempfile
import shutil
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from collections import defaultdict
import streamlit as st

# --- Streamlit UI 配置 ---
st.set_page_config(page_title="商标申请请款单生成器", layout="wide")
st.title("📝 商标申请请款单与发票申请表生成器")

# --- 辅助函数 ---

def number_to_upper(amount):
    """金额转大写（支持万、千等单位）"""
    CN_NUM = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    CN_UNIT = ['元', '拾', '佰', '仟', '万', '拾万', '佰万', '仟万', '亿']
    s = str(int(amount))[::-1]
    result = []
    for i in range(len(s)):
        digit = int(s[i])
        unit = CN_UNIT[i] if i < len(CN_UNIT) else ''
        if digit != 0:
            result.append(f"{CN_NUM[digit]}{unit}")
        else:
            if i == 0 and not result:
                result.append("零")
    formatted = ''.join(reversed(result))
    return formatted + "元整"

def create_word_doc(data, agent_fee, categories, template_path, output_path):
    """生成Word请款单"""
    try:
        # 确保模板文件存在
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件未找到: {template_path}")
        doc = Document(template_path)
    except Exception as e:
        raise Exception(f"无法加载Word模板 '{template_path}': {e}")

    num_items = len(categories)
    total_official = num_items * 270
    total_agent = num_items * agent_fee
    total_subtotal = total_official + total_agent
    total_upper = number_to_upper(total_subtotal)
    
    # 替换段落占位符
    for para in doc.paragraphs:
        if "{申请人}" in para.text:
            para.text = para.text.replace("{申请人}", data["申请人"])
        if "{事宜类型}" in para.text:
            para.text = para.text.replace("{事宜类型}", "商标注册申请")
        if "{日期}" in para.text:
            para.text = para.text.replace("{日期}", data["日期"])
        if "合计：" in para.text:
            para.text = para.text.replace("{总官费}", str(total_official))
            para.text = para.text.replace("{总代理费}", str(total_agent))
            para.text = para.text.replace("{总计}", str(total_subtotal))
            para.text = para.text.replace("{大写}", total_upper)
            
    # 处理表格
    table = doc.tables[0]
    # 移除模板中的示例行（如果存在）
    if len(table.rows) > 2:
        row_to_delete = table.rows[1]
        tbl = row_to_delete._element
        tbl.getparent().remove(tbl)

    # 添加商标信息行
    for idx, item in enumerate(categories, 1):
        row = table.add_row().cells
        row[0].text = str(idx)  # 序号
        row[1].text = item['商标名称']  # 商标名称
        row[2].text = "商标注册申请"  # 事宜
        row[3].text = item['类别']  # 类别
        row[4].text = f"¥{270}"  # 官费
        row[5].text = f"¥{agent_fee}"  # 代理费
        row[6].text = f"¥{270 + agent_fee}"  # 小计

    # 添加合计行
    total_row = table.add_row().cells
    total_row[0].merge(total_row[3])  # 合并前四个单元格 (序号, 事宜, 商标名称, 类别)
    total_row[0].text = "合计"
    # 设置合计单元格内容居中
    total_row[0].paragraphs[0].alignment = 1  # 1 代表居中对齐
    total_row[4].text = f"¥{total_official}"  # 总官费
    total_row[5].text = f"¥{total_agent}"  # 总代理费
    total_row[6].text = f"¥{total_subtotal}"  # 总计

    # 更新文件名以包含申请人和日期
    filename = f"请款单（{data['申请人']}-商标注册申请-{total_subtotal}-{data['日期']}）.docx"
    full_output_path = os.path.join(output_path, filename)
    try:
        doc.save(full_output_path)
        return full_output_path
    except Exception as e:
        raise Exception(f"保存Word文件 '{full_output_path}' 失败: {e}")

# --- 核心处理函数 ---

def extract_pdf_data_streamlit(uploaded_file):
    """从Streamlit UploadedFile对象提取数据，模拟顺序阅读。"""
    applicant = "N/A"
    unified_credit_code = "N/A"
    final_date = "N/A"
    trademarks_with_categories = []
    pending_categories = []

    # 使用 tempfile 安全地处理上传的文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    try:
        with pdfplumber.open(tmp_file_path) as pdf:
            # 提取所有页面文本并清理常见特殊空格
            all_texts = [page.extract_text().replace("\u3000", " ").replace("\xa0", " ").strip() for page in pdf.pages]
            # 使用分隔符合并文本以便于按页分割
            all_text_combined = "\n---PAGE_BREAK---\n".join(all_texts)
        pages = all_text_combined.split("\n---PAGE_BREAK---\n")

        for page_num, page_text in enumerate(pages):
            # --- 第一页：提取申请人和统一社会信用代码 ---
            if page_num == 0:
                applicant_match = re.search(r"申请人名称\(中文\)：\s*(.*?)\s*\(\s*英文\)", page_text)
                applicant = applicant_match.group(1).strip() if applicant_match else "N/A"

                # --- 修改点：支持包含字母的统一社会信用代码 ---
                unified_credit_code_match = re.search(r"统一社会信用代码：\s*([0-9A-Z]+)", page_text)
                unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
                # --- 修改点结束 ---

                # 尝试从第一页提取日期作为后备
                if final_date == "N/A":
                    date_match = re.search(r"(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)", page_text)
                    final_date = date_match.group(1).replace(" ", "") if date_match else "N/A"
                continue # 处理完第一页，继续下一页

            # --- 后续页面：提取类别或商标名 ---
            
            # 检查是否包含类别信息
            if re.search(r'类别：\d+', page_text):
                categories_found = re.findall(r'类别：(\d+)', page_text)
                pending_categories.extend(categories_found)

            # 检查是否包含委托书
            elif '商标代理委托书' in page_text:
                # --- 修改点：调整提取商标名称的正则表达式 ---
                # 新正则：更准确地定位到 '商标' 字，适应 "代理 XXX商标的如下" 的结构
                tm_name_match = re.search(r'商标代理委托书.*?代理\s+(.*?)商标\s*的\s*如下.*?事宜', page_text, re.DOTALL)
                
                tm_name = tm_name_match.group(1).strip() if tm_name_match else ""

                # 备选方案：如果上面的正则没找到，尝试更宽松的匹配
                if not tm_name:
                    fallback_match = re.search(r'代理\s+(.*?)\s*商标', page_text)
                    if fallback_match:
                        tm_name = fallback_match.group(1).strip()

                if not tm_name:
                    st.warning(f"警告：在上传的文件 '{uploaded_file.name}' 的第 {page_num + 1} 页委托书中未找到商标名称。")
                    pending_categories.clear() # 清空未关联的类别
                    continue
                # --- 修改点结束 ---

                # 提取委托书日期
                date_match = re.search(r"(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)", page_text)
                if date_match:
                    final_date = date_match.group(1).replace(" ", "") # 更新为最新的委托书日期

                # 关联：将暂存的类别与当前商标名组合
                if pending_categories:
                    for category in pending_categories:
                        trademarks_with_categories.append({
                            "商标名称": tm_name,
                            "类别": category
                        })
                    pending_categories.clear() # 关联后清空暂存区
                else:
                    # 特殊情况：委托书页面没有前面的类别信息
                    trademarks_with_categories.append({
                        "商标名称": tm_name,
                        "类别": "MANUAL_INPUT_REQUIRED" # 特殊标记
                    })
                    st.info(f"提示：上传的文件 '{uploaded_file.name}' 中的商标 '{tm_name}' 未找到自动关联的类别。")

        # --- 处理结束后，检查是否还有未关联的类别 ---
        if pending_categories:
            st.warning(f"警告：上传的文件 '{uploaded_file.name}' 处理完毕，但仍有未关联的类别 {pending_categories}。这些类别将被忽略。")

    finally:
        # 清理临时文件
        os.unlink(tmp_file_path)

    return {
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "日期": final_date,
        "商标列表": trademarks_with_categories,
        "事宜类型": "商标注册申请"
    }

def process_uploaded_files(uploaded_files, template_word_path, template_excel_path, output_dir, default_agent_fee):
    """
    处理上传的文件列表。
    :param uploaded_files: Streamlit UploadedFile 对象列表
    :param template_word_path: Word模板文件路径
    :param template_excel_path: Excel模板文件路径
    :param output_dir: 输出文件目录
    :param default_agent_fee: 默认代理费
    :return: 处理结果字典 {'success': bool, 'word_count': int, 'errors': list, ...}
    """
    results = {
        'success': False,
        'word_count': 0,
        'errors': [],
    }

    # 检查模板文件是否存在
    if not os.path.exists(template_word_path):
        error_msg = f"找不到Word模板文件: {template_word_path}"
        results['errors'].append(error_msg)
        st.error(error_msg)
        return results
    if not os.path.exists(template_excel_path):
        error_msg = f"找不到Excel模板文件: {template_excel_path}"
        results['errors'].append(error_msg)
        st.error(error_msg)
        return results

    # --- 按申请人分组数据 ---
    applicant_data_groups = defaultdict(list)
    success_count = 0

    # 第一步：遍历所有上传的PDF，提取数据并按申请人分组
    for uploaded_file in uploaded_files:
        try:
            data = extract_pdf_data_streamlit(uploaded_file)
            applicant = data["申请人"]
            applicant_data_groups[applicant].append(data)
            success_count += 1
        except Exception as e:
            error_msg = f"提取数据失败 '{uploaded_file.name}': {str(e)}"
            results['errors'].append(error_msg)
            # 在UI上显示错误，但不中断整个流程
            st.error(error_msg)

    # 第二步：为每个申请人组生成一个请款单
    generated_word_count = 0
    all_applicants_summary = []

    for applicant, data_list in applicant_data_groups.items():
        try:
            merged_trademarks = []
            latest_date = "N/A"
            unified_credit_code = "N/A"

            # --- 处理需要手动输入类别的商标 ---
            manual_input_needed = False
            manual_input_data = []

            for data in data_list:
                for tm_item in data["商标列表"]:
                    if tm_item["类别"] == "MANUAL_INPUT_REQUIRED":
                        manual_input_needed = True
                        manual_input_data.append(tm_item["商标名称"])
                    else:
                        merged_trademarks.append(tm_item)

                # 更新最新日期和统一社会信用代码
                if data["日期"] != "N/A":
                     latest_date = data["日期"]
                if data["统一社会信用代码"] != "N/A":
                     unified_credit_code = data["统一社会信用代码"]

            # --- 如果需要手动输入类别 ---
            if manual_input_needed:
                warning_msg = f"申请人 '{applicant}' 有商标需要手动输入类别: {', '.join(manual_input_data)}。这些商标已被跳过处理。"
                results['errors'].append(warning_msg)
                st.warning(warning_msg)

            # --- 准备合并后的数据结构 ---
            merged_data = {
                "申请人": applicant,
                "统一社会信用代码": unified_credit_code,
                "日期": latest_date,
                "商标列表": merged_trademarks, # 只包含自动提取的商标
                "事宜类型": "商标注册申请"
            }

            # 如果没有有效的商标项目，跳过生成请款单
            categories = merged_trademarks
            num_items = len(categories)
            if num_items == 0:
                warning_msg = f"警告：申请人 '{applicant}' 没有有效的商标项目，跳过生成请款单。"
                results['errors'].append(warning_msg)
                st.warning(warning_msg)
                continue

            # 获取代理费 (使用传入的默认值)
            agent_fee = default_agent_fee

            # 生成Word
            word_file_path = create_word_doc(merged_data, agent_fee, categories, template_word_path, output_dir)
            generated_word_count += 1

            # 收集数据到汇总列表 (为Excel)
            total_official = num_items * 270
            total_agent = num_items * agent_fee
            total_subtotal = total_official + total_agent
            all_applicants_summary.append({
                "申请人": applicant,
                "统一社会信用代码": unified_credit_code,
                "日期": latest_date,
                "总官费": total_official,
                "总代理费": total_agent,
                "总计": total_subtotal
            })

        except Exception as e:
             error_msg = f"为申请人 '{applicant}' 生成请款单时发生错误: {str(e)}"
             results['errors'].append(error_msg)
             st.error(error_msg)


    # 第三步：生成Excel汇总文件
    if all_applicants_summary:
        try:
            wb = load_workbook(template_excel_path)
            ws = wb.active
            row_num = 2
            for applicant_data in all_applicants_summary:
                # 官费行
                ws[f'B{row_num}'] = applicant_data["申请人"]
                ws[f'C{row_num}'] = applicant_data["统一社会信用代码"]
                ws[f'G{row_num}'] = applicant_data["总官费"]
                ws[f'H{row_num}'] = applicant_data["总官费"]
                ws[f'I{row_num}'] = applicant_data["总计"]
                ws[f'Q{row_num}'] = applicant_data["日期"]
                row_num += 1
                # 代理费行
                ws[f'B{row_num}'] = applicant_data["申请人"]
                ws[f'C{row_num}'] = applicant_data["统一社会信用代码"]
                ws[f'G{row_num}'] = applicant_data["总代理费"]
                ws[f'H{row_num}'] = applicant_data["总代理费"]
                ws[f'I{row_num}'] = applicant_data["总计"]
                ws[f'Q{row_num}'] = applicant_data["日期"]
                row_num += 1

            summary_date = all_applicants_summary[0]["日期"] if all_applicants_summary else "N/A"
            excel_filename = f"发票申请表-{summary_date}.xlsx"
            excel_file_path = os.path.join(output_dir, excel_filename)
            wb.save(excel_file_path)
            
        except Exception as e:
             error_msg = f"生成Excel汇总文件时发生错误: {str(e)}"
             results['errors'].append(error_msg)
             st.error(error_msg)
    else:
        warning_msg = "\n没有有效的申请人数据用于生成Excel汇总文件。"
        results['errors'].append(warning_msg)
        st.warning(warning_msg)

    results['success'] = True # 基本流程完成
    results['word_count'] = generated_word_count
    return results

# --- Streamlit 应用主逻辑 ---

# 1. 文件上传
uploaded_files = st.file_uploader("📁 请选择PDF文件", type=['pdf'], accept_multiple_files=True)

# 2. 代理费输入
default_agent_fee = st.number_input("💰 请输入默认代理费（元/项）", min_value=0, value=1000, step=100)

# 3. 处理按钮
if st.button("🚀 开始处理"):
    if not uploaded_files:
        st.warning("⚠️ 请至少上传一个PDF文件。")
    else:
        # --- 处理逻辑 ---
        with st.spinner("⏳ 正在处理文件..."):
            try:
                # 创建临时输出目录
                OUTPUT_DIR = tempfile.mkdtemp()

                # 调用核心处理函数
                # 模板文件应与 app.py 在同一目录
                processing_results = process_uploaded_files(
                    uploaded_files=uploaded_files,
                    template_word_path="请款单.docx",
                    template_excel_path="发票申请表.xlsx",
                    output_dir=OUTPUT_DIR,
                    default_agent_fee=default_agent_fee
                )

                # --- 显示处理结果 ---
                if processing_results["success"]:
                    st.success(f"✅ 处理完成！成功生成了 {processing_results['word_count']} 个请款单和 1 个发票申请表。")

                    # --- 提供文件下载 ---
                    # 打包所有生成的文件为一个 ZIP
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for root, dirs, files in os.walk(OUTPUT_DIR):
                            for file in files:
                                file_path = os.path.join(root, file)
                                # 将文件添加到 ZIP，保持相对路径
                                arcname = os.path.relpath(file_path, OUTPUT_DIR)
                                zip_file.write(file_path, arcname)

                    zip_buffer.seek(0) # 重置缓冲区指针

                    st.download_button(
                        label="📥 下载所有生成的文件 (ZIP)",
                        data=zip_buffer,
                        file_name="generated_files.zip",
                        mime="application/zip"
                    )

                else:
                    st.error("❌ 处理过程中发生错误。")
                    if processing_results.get("errors"):
                        st.text_area("🔍 错误详情:", value="\n".join(processing_results["errors"]), height=200)

            except Exception as e:
                st.error(f"💥 应用运行时发生未预期的错误: {e}")
            finally:
                # 清理临时输出目录
                if 'OUTPUT_DIR' in locals() and os.path.exists(OUTPUT_DIR):
                    shutil.rmtree(OUTPUT_DIR)

# --- README 或说明信息 ---
st.markdown("---")
st.markdown("""
### 📝 使用说明
1.  将 `请款单.docx` 和 `发票申请表.xlsx` 模板文件与 `app.py` 放在同一目录下。
2.  在左侧上传一个或多个PDF商标申请文件。
3.  输入默认代理费。
4.  点击“🚀 开始处理”。
5.  处理完成后，点击“📥 下载所有生成的文件 (ZIP)”获取结果。
""")

